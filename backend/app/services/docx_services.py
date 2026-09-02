import io

from docx import Document

from app.schemas.document import DocumentPayload


class DOCXService:
    """
    Extract structured textual content from
    Microsoft Word DOCX documents.

    Supports:
    - Paragraph extraction
    - Table extraction
    - Basic document metadata
    """

    @staticmethod
    def extract_docx(
        file_bytes: bytes,
        filename: str,
    ) -> DocumentPayload:
        """
        Extract normalized text and metadata
        from a DOCX document.
        """

        try:
            document = Document(
                io.BytesIO(file_bytes)
            )

        except Exception as exc:
            raise ValueError(
                "Unable to read DOCX document."
            ) from exc

        sections = []

        paragraph_count = 0
        table_count = len(document.tables)
        table_row_count = 0

        # =====================================================
        # PARAGRAPHS
        # =====================================================

        paragraphs = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if not text:
                continue

            paragraphs.append(text)

            paragraph_count += 1

        if paragraphs:

            sections.append(
                "Document Text\n\n"
                + "\n\n".join(paragraphs)
            )

        # =====================================================
        # TABLES
        # =====================================================

        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):

            rows = []

            for row in table.rows:

                values = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                # Ignore completely empty rows.
                if not any(values):
                    continue

                rows.append(
                    " | ".join(values)
                )

                table_row_count += 1

            if rows:

                table_section = (
                    f"Table {table_index}\n\n"
                    + "\n".join(rows)
                )

                sections.append(
                    table_section
                )

        # =====================================================
        # NORMALIZED TEXT
        # =====================================================

        extracted_text = (
            "\n\n---\n\n".join(sections)
        )

        if not extracted_text.strip():
            raise ValueError(
                "DOCX document contains no usable text."
            )

        # =====================================================
        # METADATA
        # =====================================================

        metadata = {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "table_row_count": table_row_count,
        }

        return DocumentPayload(
            filename=filename,
            file_type="docx",
            text=extracted_text,
            metadata=metadata,
        )